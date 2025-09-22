#!/usr/bin/env python3
"""
Video Chunk Transcriber

- Transcribes all chunks in an output chunk folder produced by video_splitter.py
- Aggregates per-chunk transcripts into a single full transcript
- Generates two documents:
  1) Agile meeting minutes (rule-based template)
  2) Transcript snapshot (overview & quick glances)

Requirements:
    pip install -r requirements.txt

Usage:
    python transcribe.py --chunks-dir ./output/test1_chunks --model small --language en

Outputs (created alongside the chunks directory):
    ./output/test1_chunks/transcripts/
        chunk_001.txt, chunk_001.json, ...
        full_transcript.txt, full_transcript.json
    ./output/test1_chunks/docs/
        meeting_minutes.md
        transcript_snapshot.md
"""

import os
import sys
import argparse
import json
import math
from pathlib import Path
from datetime import datetime
import shutil

# Transcription backend (local, fast, CPU/GPU):
# faster-whisper will download models automatically on first use
try:
    from faster_whisper import WhisperModel
except ImportError:
    print("Missing dependency: faster-whisper.\nPlease run: pip install faster-whisper")
    sys.exit(1)

# Optional: OpenAI for LLM refinement
try:
    import openai  # openai>=1
except Exception:
    openai = None

# Optional: DOCX/PDF export
try:
    from docx import Document  # python-docx
except Exception:
    Document = None
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
except Exception:
    A4 = None
    canvas = None
    cm = None

# Optional: diarization via pyannote (requires HF token env HUGGINGFACE_TOKEN)
try:
    from pyannote.audio import Pipeline as DiarizationPipeline  # heavy
except Exception:
    DiarizationPipeline = None

SUPPORTED_EXTS = {'.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg'}


def ensure_ffmpeg_available() -> bool:
    """Return True if ffmpeg binary is available; attempt to auto-detect winget install if missing."""
    if shutil.which('ffmpeg') and shutil.which('ffprobe'):
        return True
    try:
        local_appdata = os.environ.get('LOCALAPPDATA')
        if local_appdata:
            winget_pkgs = Path(local_appdata) / 'Microsoft' / 'WinGet' / 'Packages'
            if winget_pkgs.exists():
                candidates = list(winget_pkgs.glob('Gyan.FFmpeg*/*/bin')) + list(winget_pkgs.glob('Gyan.FFmpeg*/bin'))
                for bin_dir in candidates:
                    ffmpeg_exe = bin_dir / 'ffmpeg.exe'
                    ffprobe_exe = bin_dir / 'ffprobe.exe'
                    if ffmpeg_exe.exists() and ffprobe_exe.exists():
                        os.environ['PATH'] = str(bin_dir) + os.pathsep + os.environ.get('PATH', '')
                        if shutil.which('ffmpeg') and shutil.which('ffprobe'):
                            return True
    except Exception:
        pass
    print("Error: FFmpeg is not installed or not found in PATH. Please install it (see README.md).")
    return False


def find_chunks(chunks_dir: Path):
    files = []
    for p in sorted(chunks_dir.rglob('*')):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            files.append(p)
    return files


def transcribe_file(model: WhisperModel, filepath: Path, language: str = None, word_timestamps: bool = False):
    segments_out = []
    # beam_size 1 greedily is fast; you can tune.
    # vad_filter helps cut silence.
    translate = False  # keep language as-is
    options = dict(
        beam_size=1,
        vad_filter=True,
        language=language,
        task='transcribe',
    )
    if word_timestamps:
        options["word_timestamps"] = True
    segments, info = model.transcribe(str(filepath), **options)
    for seg in segments:
        item = {
            'start': seg.start,
            'end': seg.end,
            'text': seg.text.strip()
        }
        if word_timestamps and hasattr(seg, 'words') and seg.words:
            item['words'] = [
                {
                    'start': w.start,
                    'end': w.end,
                    'text': (w.word or '').strip()
                } for w in seg.words if w is not None
            ]
        segments_out.append(item)
    return {
        'language': info.language,
        'duration': info.duration,
        'segments': segments_out,
    }


def write_chunk_outputs(base_out: Path, chunk_index: int, data: dict):
    base_out.mkdir(parents=True, exist_ok=True)
    txt_path = base_out / f"chunk_{chunk_index:03d}.txt"
    json_path = base_out / f"chunk_{chunk_index:03d}.json"
    # TXT
    with txt_path.open('w', encoding='utf-8') as f:
        for seg in data['segments']:
            f.write(f"[{seg['start']:.2f}-{seg['end']:.2f}] {seg['text']}\n")
    # JSON
    with json_path.open('w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def aggregate_transcripts(per_chunk_data: list):
    full_text_lines = []
    full_segments = []
    for ci, item in enumerate(per_chunk_data, start=1):
        for seg in item['segments']:
            full_text_lines.append(seg['text'])
            full_segments.append({
                'chunk_index': ci,
                'start': seg['start'],
                'end': seg['end'],
                'text': seg['text'],
                'words': seg.get('words') if isinstance(seg, dict) else None,
            })
    return '\n'.join(full_text_lines), full_segments


def write_aggregate_outputs(transcripts_dir: Path, full_text: str, full_segments: list):
    txt_path = transcripts_dir / 'full_transcript.txt'
    json_path = transcripts_dir / 'full_transcript.json'
    with txt_path.open('w', encoding='utf-8') as f:
        f.write(full_text.strip() + '\n')
    with json_path.open('w', encoding='utf-8') as f:
        json.dump({'segments': full_segments}, f, ensure_ascii=False, indent=2)


# Simple rule-based Agile meeting minutes generator
# If the transcript includes explicit markers (e.g., "Action:", "Decision:"), we pick them up.
# Otherwise we provide generic sections with extracted key sentences (naive heuristics).

def extract_key_sentences(full_text: str, max_sentences: int = 10):
    # Naive split by period; pick first N non-empty sentences as a simple snapshot
    sentences = [s.strip() for s in full_text.replace('\n', ' ').split('.') if s.strip()]
    return sentences[:max_sentences]


def generate_meeting_minutes(full_text: str, meeting_title: str, facilitator: str = 'N/A', attendees: str = 'N/A') -> str:
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    sentences = extract_key_sentences(full_text, max_sentences=12)

    # Heuristic extraction
    action_items = [s for s in sentences if any(k in s.lower() for k in ['action', 'next step', 'follow up', 'assign'])]
    decisions = [s for s in sentences if any(k in s.lower() for k in ['decided', 'decision', 'agree', 'approved'])]
    risks = [s for s in sentences if any(k in s.lower() for k in ['risk', 'issue', 'concern', 'block'])]
    blockers = [s for s in sentences if any(k in s.lower() for k in ['blocked', 'blocker', 'dependency'])]

    def section(label, items):
        if not items:
            return f"- None captured"
        return '\n'.join(f"- {it}" for it in items)

    yaml_header = (
        f"---\n"
        f"title: {meeting_title}\n"
        f"datetime: {now}\n"
        f"facilitator: {facilitator}\n"
        f"attendees: {attendees}\n"
        f"---\n\n"
    )

    doc = yaml_header + f"""# Agile Meeting Minutes: {meeting_title}

- **Date/Time**: {now}
- **Facilitator**: {facilitator}
- **Attendees**: {attendees}

## Agenda
- Overview and discussion points inferred from transcript

## Key Discussion Points
{section('Discussion', sentences)}

## Decisions
{section('Decisions', decisions)}

## Action Items
{section('Actions', action_items)}

## Risks/Issues
{section('Risks', risks)}

## Blockers/Dependencies
{section('Blockers', blockers)}

## Notes
- Generated automatically from transcript using simple heuristics. Consider refining manually for accuracy.
"""
    return doc


def paragraphs_from_segments(segments: list, max_gap_s: float = 3.0, max_len_chars: int = 600):
    """Group segments into paragraphs with start/end timestamps.
    - New paragraph if time gap between segments exceeds max_gap_s, or accumulating text exceeds max_len_chars.
    """
    paragraphs = []
    cur = {"start": None, "end": None, "text": []}
    prev_end = None
    cur_len = 0
    for seg in segments:
        start = seg['start']
        end = seg['end']
        text = seg['text']
        if cur["start"] is None:
            cur["start"] = start
        if prev_end is not None and (start - prev_end > max_gap_s or cur_len + len(text) > max_len_chars):
            cur["end"] = prev_end
            paragraphs.append({"start": cur["start"], "end": cur["end"], "text": " ".join(cur["text"]).strip()})
            cur = {"start": start, "end": None, "text": []}
            cur_len = 0
        cur["text"].append(text)
        cur_len += len(text)
        prev_end = end
    if cur["text"]:
        cur["end"] = prev_end
        paragraphs.append({"start": cur["start"], "end": cur["end"], "text": " ".join(cur["text"]).strip()})
    return paragraphs


def export_docx(path: Path, title: str, sections: list):
    if Document is None:
        return False
    doc = Document()
    doc.add_heading(title, level=1)
    for sec_title, content in sections:
        doc.add_heading(sec_title, level=2)
        for line in content.split('\n'):
            doc.add_paragraph(line)
    doc.save(str(path))
    return True


def export_pdf(path: Path, title: str, sections: list):
    if canvas is None or A4 is None:
        return False
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    x_margin, y_margin = 2*cm, 2*cm
    y = height - y_margin
    c.setFont("Helvetica-Bold", 16)
    c.drawString(x_margin, y, title)
    y -= 1.0*cm
    c.setFont("Helvetica", 11)
    for sec_title, content in sections:
        c.setFont("Helvetica-Bold", 13)
        c.drawString(x_margin, y, sec_title)
        y -= 0.6*cm
        c.setFont("Helvetica", 11)
        for line in content.split('\n'):
            if y < y_margin:
                c.showPage()
                y = height - y_margin
                c.setFont("Helvetica", 11)
            c.drawString(x_margin, y, line[:110])
            y -= 0.45*cm
        y -= 0.5*cm
    c.save()
    return True


def maybe_llm_refine(minutes_md: str, snapshot_md: str, api_model: str = "gpt-4o-mini"):
    if openai is None:
        return minutes_md, snapshot_md
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return minutes_md, snapshot_md
    try:
        openai.api_key = api_key
        # Chat Completions API (new SDKs use client; keep it simple here)
        prompt_minutes = (
            "Refine the following Agile meeting minutes for clarity and concision, keeping structure intact.\n\n" + minutes_md
        )
        prompt_snapshot = (
            "Refine the following transcript snapshot for clarity and concision.\n\n" + snapshot_md
        )
        # Use legacy-style API for simplicity
        resp1 = openai.ChatCompletion.create(model=api_model, messages=[{"role": "user", "content": prompt_minutes}])
        resp2 = openai.ChatCompletion.create(model=api_model, messages=[{"role": "user", "content": prompt_snapshot}])
        m1 = resp1.choices[0].message["content"]
        m2 = resp2.choices[0].message["content"]
        return m1, m2
    except Exception:
        return minutes_md, snapshot_md


def generate_snapshot(full_text: str, meeting_title: str) -> str:
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    sentences = extract_key_sentences(full_text, max_sentences=8)
    preview = '\n'.join(f"- {s}" for s in sentences)
    words = len(full_text.split())
    chars = len(full_text)
    doc = f"""# Transcript Snapshot: {meeting_title}

- **Generated**: {now}
- **Approx. length**: {words} words, {chars} chars

## Quick Glance
{preview}

## Tip
- For best results, mark explicit items in meetings: "Decision:", "Action:", "Risk:", etc. The generator will pick them up.
"""
    return doc


def main():
    parser = argparse.ArgumentParser(description="Transcribe video chunks and generate documents")
    parser.add_argument('--chunks-dir', required=True, help='Path to chunks directory, e.g., ./output/test1_chunks')
    parser.add_argument('--model', default='small', help='faster-whisper model size (tiny, base, small, medium, large)')
    parser.add_argument('--device', default='auto', choices=['auto', 'cpu', 'cuda'], help='Device to run on')
    parser.add_argument('--compute-type', default='auto', help='Compute type: auto (default), int8, int8_float16, float16, float32')
    parser.add_argument('--language', default=None, help='Language code, e.g., en, hi, etc. Leave None for auto')
    parser.add_argument('--word-timestamps', action='store_true', help='Include word-level timestamps in output')
    parser.add_argument('--para-gap', type=float, default=3.0, help='Max gap (s) between segments before starting new paragraph')
    parser.add_argument('--para-max-chars', type=int, default=600, help='Max paragraph text length before splitting')
    parser.add_argument('--facilitator', default='N/A', help='Facilitator name for YAML header')
    parser.add_argument('--attendees', default='N/A', help='Comma-separated attendees for YAML header')
    parser.add_argument('--use-llm', action='store_true', help='Refine minutes and snapshot with OpenAI if OPENAI_API_KEY is set')
    parser.add_argument('--diarize', action='store_true', help='Attempt speaker diarization with pyannote (requires HF token)')
    args = parser.parse_args()

    chunks_dir = Path(args.chunks_dir).resolve()
    if not chunks_dir.exists():
        print(f"Error: chunks directory not found: {chunks_dir}")
        sys.exit(1)

    if not ensure_ffmpeg_available():
        sys.exit(1)

    chunk_files = find_chunks(chunks_dir)
    if not chunk_files:
        print(f"No chunk files found under: {chunks_dir}")
        sys.exit(1)

    # Outputs
    transcripts_dir = chunks_dir / 'transcripts'
    docs_dir = chunks_dir / 'docs'
    transcripts_dir.mkdir(parents=True, exist_ok=True)
    docs_dir.mkdir(parents=True, exist_ok=True)

    # Initialize model
    device = None if args.device == 'auto' else args.device
    selected_device = device or 'cpu'

    # Auto-select compute type if requested
    compute_type = args.compute_type
    if compute_type == 'auto':
        if selected_device == 'cuda':
            compute_type = 'float16'  # broadly supported on CUDA
        else:
            compute_type = 'int8'     # efficient on CPU

    # Try initialization with graceful fallbacks
    try:
        model = WhisperModel(args.model, device=selected_device, compute_type=compute_type)
    except ValueError:
        # Fallback chain for incompatibilities
        fallback_order = []
        if selected_device == 'cuda':
            fallback_order = ['float16', 'int8_float16', 'float32']
        else:
            fallback_order = ['int8', 'float32']
        last_err = None
        for ct in fallback_order:
            try:
                model = WhisperModel(args.model, device=selected_device, compute_type=ct)
                compute_type = ct
                break
            except Exception as e:
                last_err = e
                continue
        else:
            print(f"Failed to initialize WhisperModel with compatible compute type. Last error: {last_err}")
            sys.exit(1)

    per_chunk_data = []
    for idx, file in enumerate(chunk_files, start=1):
        print(f"Transcribing chunk {idx}/{len(chunk_files)}: {file.name}")
        data = transcribe_file(model, file, language=args.language, word_timestamps=args.word_timestamps)
        per_chunk_data.append(data)
        write_chunk_outputs(transcripts_dir, idx, data)

    # Aggregate
    full_text, full_segments = aggregate_transcripts(per_chunk_data)
    write_aggregate_outputs(transcripts_dir, full_text, full_segments)

    # Paragraphs with timestamps for snapshot enrichment
    paragraphs = paragraphs_from_segments(full_segments, max_gap_s=args.para_gap, max_len_chars=args.para_max_chars)
    para_overview = '\n'.join(
        f"[{p['start']:.2f}-{p['end']:.2f}] {p['text']}" for p in paragraphs[:20]
    )

    # Docs
    meeting_title = chunks_dir.name.replace('_', ' ')
    minutes_md = generate_meeting_minutes(full_text, meeting_title, facilitator=args.facilitator, attendees=args.attendees)
    snapshot_md = generate_snapshot(full_text, meeting_title) + "\n\n## Paragraph Overview (first 20)\n" + para_overview

    # Optional LLM refinement
    if args.use_llm:
        minutes_md, snapshot_md = maybe_llm_refine(minutes_md, snapshot_md)

    (docs_dir / 'meeting_minutes.md').write_text(minutes_md, encoding='utf-8')
    (docs_dir / 'transcript_snapshot.md').write_text(snapshot_md, encoding='utf-8')

    # Export to DOCX / PDF if libraries available
    export_docx(docs_dir / 'meeting_minutes.docx', meeting_title + ' - Minutes', [("Meeting Minutes", minutes_md)])
    export_docx(docs_dir / 'transcript_snapshot.docx', meeting_title + ' - Snapshot', [("Transcript Snapshot", snapshot_md)])
    export_pdf(docs_dir / 'meeting_minutes.pdf', meeting_title + ' - Minutes', [("Meeting Minutes", minutes_md)])
    export_pdf(docs_dir / 'transcript_snapshot.pdf', meeting_title + ' - Snapshot', [("Transcript Snapshot", snapshot_md)])

    print("\nTranscription complete.")
    print(f"Per-chunk transcripts: {transcripts_dir}")
    print(f"Full transcript: {transcripts_dir / 'full_transcript.txt'}")
    print(f"Documents: {docs_dir}")


if __name__ == '__main__':
    main()
