#!/usr/bin/env python3
"""
File Management Microservice
Handles file operations, dated folder creation, and output organization.
"""

import os
import sys
import json
import shutil
import logging
import zipfile
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, Optional

try:
    from flask import Flask, request, jsonify
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileManagementService:
    def __init__(self, base_output_dir: str = "./output", input_dir: str = "./input"):
        self.base_output_dir = Path(base_output_dir)
        self.input_dir = Path(input_dir)
        self.base_output_dir.mkdir(exist_ok=True)
        self.input_dir.mkdir(exist_ok=True)
        
        # Start background cleanup thread
        self.cleanup_thread = threading.Thread(target=self._cleanup_old_files, daemon=True)
        self.cleanup_thread.start()
        
        # Track processed files for cleanup
        self.processed_files = set()
    
    def create_dated_folder(self, meeting_title: str, meeting_date: Optional[datetime] = None) -> str:
        """
        Create a dated folder for meeting outputs.
        
        Args:
            meeting_title: Title of the meeting
            meeting_date: Date of the meeting (default: today)
            
        Returns:
            Path to the created folder
        """
        if meeting_date is None:
            meeting_date = datetime.now()
        
        # Create folder name: YYYY-MM-DD_MeetingTitle
        safe_title = "".join(c for c in meeting_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        folder_name = f"{meeting_date.strftime('%Y-%m-%d')}_{safe_title}"
        
        folder_path = self.base_output_dir / folder_name
        folder_path.mkdir(exist_ok=True)
        
        logger.info(f"Created dated folder: {folder_path}")
        return str(folder_path)
    
    def save_transcript(self, transcript_text: str, output_folder: str, 
                       filename: str = "transcript.txt") -> Dict[str, Any]:
        """
        Save transcript to a text file.
        
        Args:
            transcript_text: The transcript content
            output_folder: Folder to save the file
            filename: Name of the file
            
        Returns:
            Dictionary with success status and file path
        """
        try:
            output_path = Path(output_folder) / filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(transcript_text)
            
            logger.info(f"Transcript saved to: {output_path}")
            
            return {
                "success": True,
                "file_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to save transcript: {e}")
            return {
                "success": False,
                "file_path": None,
                "file_size": 0,
                "error": str(e)
            }
    
    def save_meeting_minutes_docx(self, meeting_data: Dict[str, Any], output_folder: str,
                                 filename: str = "meeting_minutes.docx") -> Dict[str, Any]:
        """
        Save meeting minutes as DOCX file.
        
        Args:
            meeting_data: The meeting minutes data
            output_folder: Folder to save the file
            filename: Name of the file
            
        Returns:
            Dictionary with success status and file path
        """
        try:
            output_path = Path(output_folder) / filename
            
            # Import here to avoid dependency issues
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                return {
                    "success": False,
                    "file_path": None,
                    "file_size": 0,
                    "error": "python-docx not available"
                }
            
            doc = Document()
            
            # Title
            title = doc.add_heading(meeting_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Meeting info
            doc.add_heading('Meeting Information', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = 'Date'
            info_table.cell(0, 1).text = meeting_data['date'].strftime('%B %d, %Y at %I:%M %p')
            info_table.cell(1, 0).text = 'Participants'
            info_table.cell(1, 1).text = ', '.join(meeting_data['speakers'])
            info_table.cell(2, 0).text = 'Meeting Type'
            info_table.cell(2, 1).text = 'Agile TPM Program Review'
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(meeting_data['tpm_summary'])
            
            # Program Metrics
            doc.add_heading('Program Metrics', level=1)
            metrics_table = doc.add_table(rows=1, cols=5)
            metrics_table.style = 'Table Grid'
            
            headers = ['Epics', 'User Stories', 'Technical Tasks', 'Risks', 'Action Items']
            for i, header in enumerate(headers):
                metrics_table.cell(0, i).text = header
            
            # Add metrics data
            metrics_row = metrics_table.add_row()
            metrics_row.cells[0].text = str(len(meeting_data['artifacts'].get('epics', [])))
            metrics_row.cells[1].text = str(len(meeting_data['artifacts'].get('user_stories', [])))
            metrics_row.cells[2].text = str(len(meeting_data['artifacts'].get('technical_tasks', [])))
            metrics_row.cells[3].text = str(len(meeting_data['artifacts'].get('risks', [])))
            metrics_row.cells[4].text = str(len(meeting_data['artifacts'].get('action_items', [])))
            
            # Epics
            if meeting_data['artifacts'].get('epics'):
                doc.add_heading('Epics & Major Initiatives', level=1)
                for epic in meeting_data['artifacts']['epics']:
                    doc.add_heading(epic['title'], level=2)
                    doc.add_paragraph(epic['description'])
                    doc.add_paragraph(f"Priority: {epic['priority']}")
            
            # Risks
            if meeting_data['artifacts'].get('risks'):
                doc.add_heading('Risks & Issues', level=1)
                for risk in meeting_data['artifacts']['risks']:
                    doc.add_heading(risk['title'], level=2)
                    doc.add_paragraph(risk['description'])
                    doc.add_paragraph(f"Impact: {risk['impact']} | Probability: {risk['probability']}")
            
            # Decisions
            if meeting_data['artifacts'].get('decisions'):
                doc.add_heading('Decisions Made', level=1)
                for decision in meeting_data['artifacts']['decisions']:
                    doc.add_heading(decision['title'], level=2)
                    doc.add_paragraph(decision['description'])
                    doc.add_paragraph(f"Rationale: {decision['rationale']}")
            
            # Action Items
            if meeting_data['artifacts'].get('action_items'):
                doc.add_heading('Action Items', level=1)
                for item in meeting_data['artifacts']['action_items']:
                    doc.add_heading(f"Action: {item['action'][:50]}...", level=2)
                    doc.add_paragraph(f"Owner: {item['owner']}")
                    doc.add_paragraph(f"Due Date: {item['due_date']}")
                    doc.add_paragraph(f"Priority: {item['priority']}")
            
            # Timeline
            if meeting_data['sprint_info'].get('dates') or meeting_data['sprint_info'].get('sprints'):
                doc.add_heading('Timeline & Milestones', level=1)
                if meeting_data['sprint_info'].get('dates'):
                    doc.add_heading('Key Dates', level=2)
                    for date in meeting_data['sprint_info']['dates']:
                        doc.add_paragraph(f"• {date}")
                if meeting_data['sprint_info'].get('sprints'):
                    doc.add_heading('Sprint References', level=2)
                    for sprint in meeting_data['sprint_info']['sprints']:
                        doc.add_paragraph(f"• {sprint}")
            
            # Save document
            doc.save(output_path)
            logger.info(f"DOCX document saved to: {output_path}")
            
            return {
                "success": True,
                "file_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to save DOCX: {e}")
            return {
                "success": False,
                "file_path": None,
                "file_size": 0,
                "error": str(e)
            }
    
    def save_meeting_minutes_html(self, meeting_data: Dict[str, Any], output_folder: str,
                                 filename: str = "meeting_minutes.html") -> Dict[str, Any]:
        """
        Save meeting minutes as HTML file.
        
        Args:
            meeting_data: The meeting minutes data
            output_folder: Folder to save the file
            filename: Name of the file
            
        Returns:
            Dictionary with success status and file path
        """
        try:
            output_path = Path(output_folder) / filename
            
            # Simple HTML generation
            artifacts = meeting_data['artifacts']
            
            html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{meeting_data['title']}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; max-width: 1000px; margin: 0 auto; padding: 20px; background-color: #f8f9fa; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
        .section {{ background: white; margin-bottom: 20px; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .section h2 {{ color: #2c3e50; border-left: 4px solid #3498db; padding-left: 15px; margin-top: 0; }}
        .epic {{ background: #e8f5e8; border-left: 4px solid #27ae60; padding: 10px; margin: 5px 0; }}
        .risk {{ background: #ffebee; border-left: 4px solid #f44336; padding: 10px; margin: 5px 0; }}
        .decision {{ background: #e0f2f1; border-left: 4px solid #009688; padding: 10px; margin: 5px 0; }}
        .action-item {{ background: #fff8e1; border-left: 4px solid #ffc107; padding: 10px; margin: 5px 0; }}
        .transcription {{ max-height: 400px; overflow-y: auto; border: 1px solid #ddd; padding: 15px; background: #fafafa; }}
        .speaker {{ font-weight: bold; color: #3498db; }}
        .timestamp {{ color: #7f8c8d; font-size: 0.9em; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{meeting_data['title']}</h1>
        <p><strong>Date:</strong> {meeting_data['date'].strftime('%B %d, %Y at %I:%M %p')}</p>
        <p><strong>Participants:</strong> {', '.join(meeting_data['speakers'])}</p>
        <p><strong>Meeting Type:</strong> Agile TPM Program Review</p>
    </div>

    <div class="section">
        <h2>📋 Executive Summary</h2>
        <p>{meeting_data['tpm_summary']}</p>
    </div>

    <div class="section">
        <h2>🎯 Epics & Major Initiatives</h2>
        {''.join([f'<div class="epic"><strong>{epic["title"]}</strong><br>{epic["description"]}<br><small>Priority: {epic["priority"]}</small></div>' for epic in artifacts.get('epics', [])])}
    </div>

    <div class="section">
        <h2>⚠️ Risks & Issues</h2>
        {''.join([f'<div class="risk"><strong>{risk["title"]}</strong><br>{risk["description"]}<br><small>Impact: {risk["impact"]} | Probability: {risk["probability"]}</small></div>' for risk in artifacts.get('risks', [])])}
    </div>

    <div class="section">
        <h2>✅ Decisions Made</h2>
        {''.join([f'<div class="decision"><strong>{decision["title"]}</strong><br>{decision["description"]}<br><small>Rationale: {decision["rationale"]}</small></div>' for decision in artifacts.get('decisions', [])])}
    </div>

    <div class="section">
        <h2>📝 Action Items</h2>
        {''.join([f'<div class="action-item"><strong>Action:</strong> {item["action"]}<br><strong>Owner:</strong> {item["owner"]}<br><strong>Due Date:</strong> {item["due_date"]}<br><small>Priority: {item["priority"]}</small></div>' for item in artifacts.get('action_items', [])])}
    </div>

    <div class="section">
        <h2>💬 Full Transcription</h2>
        <div class="transcription">
            {''.join([f'<p><span class="speaker">{item["speaker"]}:</span> {item["content"]}</p>' for item in meeting_data['transcription']])}
        </div>
    </div>
</body>
</html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML document saved to: {output_path}")
            
            return {
                "success": True,
                "file_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to save HTML: {e}")
            return {
                "success": False,
                "file_path": None,
                "file_size": 0,
                "error": str(e)
            }
    
    def copy_video_to_output(self, video_path: str, output_folder: str,
                           filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Copy the original video file to the output folder.
        
        Args:
            video_path: Path to the original video file
            output_folder: Folder to copy the video to
            filename: New filename (default: original filename)
            
        Returns:
            Dictionary with success status and file path
        """
        try:
            source_path = Path(video_path)
            if not source_path.exists():
                return {
                    "success": False,
                    "file_path": None,
                    "file_size": 0,
                    "error": f"Source video not found: {video_path}"
                }
            
            if filename is None:
                filename = source_path.name
            
            dest_path = Path(output_folder) / filename
            
            shutil.copy2(source_path, dest_path)
            
            logger.info(f"Video copied to: {dest_path}")
            
            return {
                "success": True,
                "file_path": str(dest_path),
                "file_size": dest_path.stat().st_size,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to copy video: {e}")
            return {
                "success": False,
                "file_path": None,
                "file_size": 0,
                "error": str(e)
            }
    
    def create_workflow_summary(self, output_folder: str, workflow_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Create a workflow summary file with all processing details.
        
        Args:
            output_folder: Folder to save the summary
            workflow_data: Data about the workflow execution
            
        Returns:
            Dictionary with success status and file path
        """
        try:
            output_path = Path(output_folder) / "workflow_summary.json"
            
            # Add timestamp
            workflow_data['processed_at'] = datetime.now().isoformat()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(workflow_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Workflow summary saved to: {output_path}")
            
            return {
                "success": True,
                "file_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to save workflow summary: {e}")
            return {
                "success": False,
                "file_path": None,
                "file_size": 0,
                "error": str(e)
            }
    
    def delete_input_file(self, video_path: str) -> Dict[str, Any]:
        """
        Delete the input video file after processing.
        
        Args:
            video_path: Path to the input video file
            
        Returns:
            Dictionary with success status
        """
        try:
            file_path = Path(video_path)
            
            # Only delete files from the input directory for safety
            if not str(file_path).startswith(str(self.input_dir)):
                return {
                    "success": False,
                    "error": "File is not in input directory, cannot delete for safety"
                }
            
            if file_path.exists():
                file_path.unlink()
                logger.info(f"Input file deleted: {video_path}")
                return {
                    "success": True,
                    "message": f"Input file deleted: {file_path.name}"
                }
            else:
                return {
                    "success": False,
                    "error": "File not found"
                }
                
        except Exception as e:
            logger.error(f"Failed to delete input file: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def zip_output_folder(self, output_folder: str, meeting_title: str) -> Dict[str, Any]:
        """
        Create a zip file containing all output files.
        
        Args:
            output_folder: Path to the output folder
            meeting_title: Title of the meeting for zip filename
            
        Returns:
            Dictionary with success status and zip file path
        """
        try:
            folder_path = Path(output_folder)
            if not folder_path.exists():
                return {
                    "success": False,
                    "error": "Output folder does not exist"
                }
            
            # Create zip filename
            safe_title = "".join(c for c in meeting_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')
            zip_filename = f"{safe_title}_output.zip"
            zip_path = folder_path / zip_filename
            
            # Create zip file
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in folder_path.rglob('*'):
                    if file_path.is_file() and file_path.name != zip_filename:
                        # Add file to zip with relative path
                        arcname = file_path.relative_to(folder_path)
                        zipf.write(file_path, arcname)
            
            logger.info(f"Output folder zipped: {zip_path}")
            
            return {
                "success": True,
                "zip_path": str(zip_path),
                "zip_size": zip_path.stat().st_size,
                "files_zipped": len([f for f in folder_path.rglob('*') if f.is_file() and f.name != zip_filename])
            }
            
        except Exception as e:
            logger.error(f"Failed to zip output folder: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _cleanup_old_files(self):
        """
        Background thread to clean up old output files after 24 hours.
        """
        while True:
            try:
                current_time = datetime.now()
                cutoff_time = current_time - timedelta(hours=24)
                
                # Clean up old output folders
                for folder_path in self.base_output_dir.iterdir():
                    if folder_path.is_dir():
                        # Check if folder is older than 24 hours
                        folder_mtime = datetime.fromtimestamp(folder_path.stat().st_mtime)
                        if folder_mtime < cutoff_time:
                            try:
                                shutil.rmtree(folder_path)
                                logger.info(f"Cleaned up old output folder: {folder_path}")
                            except Exception as e:
                                logger.error(f"Failed to clean up folder {folder_path}: {e}")
                
                # Sleep for 1 hour before next cleanup
                time.sleep(3600)
                
            except Exception as e:
                logger.error(f"Error in cleanup thread: {e}")
                time.sleep(3600)  # Sleep for 1 hour on error
    
    def get_cleanup_status(self) -> Dict[str, Any]:
        """
        Get status of file cleanup operations.
        
        Returns:
            Dictionary with cleanup status information
        """
        try:
            current_time = datetime.now()
            cutoff_time = current_time - timedelta(hours=24)
            
            total_folders = 0
            old_folders = 0
            total_size = 0
            
            for folder_path in self.base_output_dir.iterdir():
                if folder_path.is_dir():
                    total_folders += 1
                    folder_mtime = datetime.fromtimestamp(folder_path.stat().st_mtime)
                    if folder_mtime < cutoff_time:
                        old_folders += 1
                    
                    # Calculate folder size
                    folder_size = sum(f.stat().st_size for f in folder_path.rglob('*') if f.is_file())
                    total_size += folder_size
            
            return {
                "success": True,
                "total_folders": total_folders,
                "old_folders": old_folders,
                "total_size_mb": round(total_size / (1024 * 1024), 2),
                "next_cleanup": "Every hour",
                "cutoff_hours": 24
            }
            
        except Exception as e:
            logger.error(f"Failed to get cleanup status: {e}")
            return {
                "success": False,
                "error": str(e)
            }


# Flask API Service
if FLASK_AVAILABLE:
    app = Flask(__name__)
    file_service = FileManagementService()

    @app.route('/health', methods=['GET'])
    def health_check():
        """Health check endpoint."""
        return jsonify({
            "status": "healthy",
            "service": "file_management",
            "base_output_dir": str(file_service.base_output_dir)
        })

    @app.route('/create-dated-folder', methods=['POST'])
    def create_dated_folder():
        """Create a dated folder for meeting outputs."""
        try:
            data = request.get_json()
            
            if not data or 'meeting_title' not in data:
                return jsonify({
                    "success": False,
                    "error": "meeting_title is required"
                }), 400
            
            meeting_title = data['meeting_title']
            meeting_date = data.get('meeting_date')
            
            if meeting_date:
                try:
                    meeting_date = datetime.fromisoformat(meeting_date)
                except ValueError:
                    meeting_date = None
            
            folder_path = file_service.create_dated_folder(meeting_title, meeting_date)
            
            return jsonify({
                "success": True,
                "folder_path": folder_path
            })
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/save-transcript', methods=['POST'])
    def save_transcript():
        """Save transcript to a text file."""
        try:
            data = request.get_json()
            
            if not data or 'transcript_text' not in data or 'output_folder' not in data:
                return jsonify({
                    "success": False,
                    "error": "transcript_text and output_folder are required"
                }), 400
            
            result = file_service.save_transcript(
                data['transcript_text'],
                data['output_folder'],
                data.get('filename', 'transcript.txt')
            )
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/save-meeting-minutes-docx', methods=['POST'])
    def save_meeting_minutes_docx():
        """Save meeting minutes as DOCX file."""
        try:
            data = request.get_json()
            
            if not data or 'meeting_data' not in data or 'output_folder' not in data:
                return jsonify({
                    "success": False,
                    "error": "meeting_data and output_folder are required"
                }), 400
            
            result = file_service.save_meeting_minutes_docx(
                data['meeting_data'],
                data['output_folder'],
                data.get('filename', 'meeting_minutes.docx')
            )
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/save-meeting-minutes-html', methods=['POST'])
    def save_meeting_minutes_html():
        """Save meeting minutes as HTML file."""
        try:
            data = request.get_json()
            
            if not data or 'meeting_data' not in data or 'output_folder' not in data:
                return jsonify({
                    "success": False,
                    "error": "meeting_data and output_folder are required"
                }), 400
            
            result = file_service.save_meeting_minutes_html(
                data['meeting_data'],
                data['output_folder'],
                data.get('filename', 'meeting_minutes.html')
            )
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/copy-video', methods=['POST'])
    def copy_video():
        """Copy video file to output folder."""
        try:
            data = request.get_json()
            
            if not data or 'video_path' not in data or 'output_folder' not in data:
                return jsonify({
                    "success": False,
                    "error": "video_path and output_folder are required"
                }), 400
            
            result = file_service.copy_video_to_output(
                data['video_path'],
                data['output_folder'],
                data.get('filename')
            )
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/create-workflow-summary', methods=['POST'])
    def create_workflow_summary():
        """Create workflow summary file."""
        try:
            data = request.get_json()
            
            if not data or 'output_folder' not in data or 'workflow_data' not in data:
                return jsonify({
                    "success": False,
                    "error": "output_folder and workflow_data are required"
                }), 400
            
            result = file_service.create_workflow_summary(
                data['output_folder'],
                data['workflow_data']
            )
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/delete-input-file', methods=['POST'])
    def delete_input_file():
        """Delete input video file after processing."""
        try:
            data = request.get_json()
            
            if not data or 'video_path' not in data:
                return jsonify({
                    "success": False,
                    "error": "video_path is required"
                }), 400
            
            result = file_service.delete_input_file(data['video_path'])
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/zip-output-folder', methods=['POST'])
    def zip_output_folder():
        """Create zip file of output folder."""
        try:
            data = request.get_json()
            
            if not data or 'output_folder' not in data or 'meeting_title' not in data:
                return jsonify({
                    "success": False,
                    "error": "output_folder and meeting_title are required"
                }), 400
            
            result = file_service.zip_output_folder(
                data['output_folder'],
                data['meeting_title']
            )
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/cleanup-status', methods=['GET'])
    def cleanup_status():
        """Get cleanup status information."""
        try:
            result = file_service.get_cleanup_status()
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    def run_service(host='localhost', port=5003, debug=False, base_output_dir='./output'):
        """Run the file management service."""
        global file_service
        file_service = FileManagementService(base_output_dir)
        
        logger.info(f"Starting File Management Service on {host}:{port}")
        app.run(host=host, port=port, debug=debug)

else:
    def run_service(host='localhost', port=5003, debug=False, base_output_dir='./output'):
        """Fallback when Flask is not available."""
        logger.error("Flask not available. Cannot run as web service.")
        sys.exit(1)


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="File Management Microservice")
    parser.add_argument('--host', default='localhost', help='Host to bind to')
    parser.add_argument('--port', type=int, default=5003, help='Port to bind to')
    parser.add_argument('--base-output-dir', default='./output', help='Base output directory')
    parser.add_argument('--debug', action='store_true', help='Enable debug mode')
    
    args = parser.parse_args()
    
    run_service(host=args.host, port=args.port, debug=args.debug, base_output_dir=args.base_output_dir)
