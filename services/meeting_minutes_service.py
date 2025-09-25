#!/usr/bin/env python3
"""
Meeting Minutes Generation Microservice
Handles Agile TPM meeting minutes generation with REST API interface.
"""

import os
import sys
import json
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from jinja2 import Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from flask import Flask, request, jsonify
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class MeetingMinutesService:
    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key
        if OPENAI_AVAILABLE and api_key:
            openai.api_key = api_key
            self.client = openai.OpenAI(api_key=api_key)
        else:
            self.client = None
    
    def parse_transcription(self, transcription_text: str) -> list:
        """Parse transcription text and extract speaker information."""
        lines = transcription_text.strip().split('\n')
        parsed_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Try to match speaker patterns
            speaker_match = re.match(r'^(\[?(\d{2}:\d{2}:\d{2})\]?\s*)?([^:]+):\s*(.*)$', line)
            if speaker_match:
                timestamp = speaker_match.group(2)
                speaker = speaker_match.group(3).strip()
                content = speaker_match.group(4).strip()
                
                parsed_content.append({
                    'timestamp': timestamp,
                    'speaker': speaker,
                    'content': content
                })
            else:
                # If no speaker pattern, treat as continuation
                if parsed_content:
                    parsed_content[-1]['content'] += ' ' + line
                else:
                    parsed_content.append({
                        'timestamp': None,
                        'speaker': 'Unknown',
                        'content': line
                    })
        
        return parsed_content

    def extract_agile_artifacts(self, parsed_content: list) -> Dict[str, Any]:
        """Extract Agile artifacts like epics, stories, tasks, etc."""
        if not self.client:
            return self._extract_agile_artifacts_simple(parsed_content)
        
        try:
            full_text = '\n'.join([f"{item['speaker']}: {item['content']}" for item in parsed_content])
            
            prompt = f"""
            Analyze this meeting transcription and extract Agile TPM artifacts in JSON format:
            
            1. Epics (major initiatives/features)
            2. User Stories (business requirements)
            3. Technical Tasks (implementation work)
            4. Risks (potential issues)
            5. Dependencies (blockers or prerequisites)
            6. Decisions (technical or business decisions made)
            7. Action Items (specific tasks with owners)
            
            Return a JSON object with these keys:
            {{
                "epics": [{{"title": "Epic name", "description": "Epic description", "priority": "High/Medium/Low"}}],
                "user_stories": [{{"title": "Story title", "description": "Story description", "acceptance_criteria": "Criteria"}}],
                "technical_tasks": [{{"title": "Task title", "description": "Task description", "effort": "Story points or time estimate"}}],
                "risks": [{{"title": "Risk title", "description": "Risk description", "impact": "High/Medium/Low", "probability": "High/Medium/Low"}}],
                "dependencies": [{{"title": "Dependency title", "description": "Dependency description", "blocker": "What it blocks"}}],
                "decisions": [{{"title": "Decision title", "description": "Decision description", "rationale": "Why this decision"}}],
                "action_items": [{{"action": "Action description", "owner": "Person responsible", "due_date": "Date or sprint", "priority": "High/Medium/Low"}}]
            }}
            
            Transcription:
            {full_text[:4000]}
            
            Return only valid JSON.
            """
            
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            
            result = response.choices[0].message.content.strip()
            try:
                return json.loads(result)
            except json.JSONDecodeError:
                return self._extract_agile_artifacts_simple(parsed_content)
                
        except Exception as e:
            logger.error(f"AI analysis failed: {e}. Using simple extraction.")
            return self._extract_agile_artifacts_simple(parsed_content)

    def _extract_agile_artifacts_simple(self, parsed_content: list) -> Dict[str, Any]:
        """Simple keyword-based Agile artifact extraction."""
        artifacts = {
            "epics": [],
            "user_stories": [],
            "technical_tasks": [],
            "risks": [],
            "dependencies": [],
            "decisions": [],
            "action_items": []
        }
        
        content_text = ' '.join([item['content'] for item in parsed_content])
        
        # Extract epics (major features/initiatives)
        epic_keywords = ['epic', 'initiative', 'feature', 'release', 'milestone']
        for keyword in epic_keywords:
            if keyword in content_text.lower():
                sentences = re.split(r'[.!?]+', content_text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        artifacts["epics"].append({
                            "title": f"Epic: {keyword.title()}",
                            "description": sentence.strip(),
                            "priority": "Medium"
                        })
                        break
        
        # Extract risks
        risk_keywords = ['risk', 'concern', 'issue', 'problem', 'challenge', 'blocker']
        for keyword in risk_keywords:
            if keyword in content_text.lower():
                sentences = re.split(r'[.!?]+', content_text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        artifacts["risks"].append({
                            "title": f"Risk: {keyword.title()}",
                            "description": sentence.strip(),
                            "impact": "Medium",
                            "probability": "Medium"
                        })
                        break
        
        # Extract action items
        action_keywords = ['action', 'todo', 'task', 'follow up', 'next step', 'will do', 'need to']
        for keyword in action_keywords:
            if keyword in content_text.lower():
                sentences = re.split(r'[.!?]+', content_text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        artifacts["action_items"].append({
                            "action": sentence.strip(),
                            "owner": "TBD",
                            "due_date": "TBD",
                            "priority": "Medium"
                        })
                        break
        
        # Extract decisions
        decision_keywords = ['decided', 'decision', 'agree', 'approved', 'concluded', 'resolved']
        for keyword in decision_keywords:
            if keyword in content_text.lower():
                sentences = re.split(r'[.!?]+', content_text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        artifacts["decisions"].append({
                            "title": f"Decision: {keyword.title()}",
                            "description": sentence.strip(),
                            "rationale": "Discussed in meeting"
                        })
                        break
        
        return artifacts

    def extract_sprint_info(self, parsed_content: list) -> Dict[str, Any]:
        """Extract sprint and timeline information."""
        content_text = ' '.join([item['content'] for item in parsed_content])
        
        # Look for dates and deadlines
        date_patterns = [
            r'(\d{1,2}(?:st|nd|rd|th)?\s+(?:of\s+)?(?:January|February|March|April|May|June|July|August|September|October|November|December))',
            r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(?:st|nd|rd|th)?)',
            r'(\d{1,2}/\d{1,2}/\d{2,4})',
            r'(\d{4}-\d{2}-\d{2})'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, content_text, re.IGNORECASE)
            dates.extend(matches)
        
        # Look for sprint references
        sprint_keywords = ['sprint', 'iteration', 'cycle', 'milestone', 'release']
        sprints = []
        for keyword in sprint_keywords:
            if keyword in content_text.lower():
                sentences = re.split(r'[.!?]+', content_text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        sprints.append(sentence.strip())
                        break
        
        return {
            "dates": list(set(dates)),
            "sprints": list(set(sprints))
        }

    def generate_tpm_summary(self, parsed_content: list, artifacts: Dict[str, Any], 
                           sprint_info: Dict[str, Any]) -> str:
        """Generate TPM-style executive summary."""
        if not self.client:
            return self._generate_tpm_summary_simple(parsed_content, artifacts, sprint_info)
        
        try:
            full_text = '\n'.join([f"{item['speaker']}: {item['content']}" for item in parsed_content])
            
            prompt = f"""
            Write an executive summary for a Technical Program Manager meeting minutes document.
            Focus on:
            1. Key program status and progress
            2. Critical milestones and deadlines
            3. Major risks and blockers
            4. Resource and timeline implications
            5. Next steps and recommendations
            
            Meeting content:
            {full_text[:3000]}
            
            Extracted artifacts: {json.dumps(artifacts, indent=2)}
            Sprint info: {json.dumps(sprint_info, indent=2)}
            
            Write 2-3 paragraphs in a professional TPM style.
            """
            
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.3
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.error(f"AI summary generation failed: {e}. Using simple summary.")
            return self._generate_tpm_summary_simple(parsed_content, artifacts, sprint_info)

    def _generate_tpm_summary_simple(self, parsed_content: list, artifacts: Dict[str, Any], 
                                   sprint_info: Dict[str, Any]) -> str:
        """Simple TPM summary generation."""
        total_actions = len(artifacts.get("action_items", []))
        total_risks = len(artifacts.get("risks", []))
        total_decisions = len(artifacts.get("decisions", []))
        
        summary = f"""
        Program Status Update: This meeting covered key program activities with {total_actions} action items identified, {total_risks} risks discussed, and {total_decisions} decisions made. 
        
        Critical milestones and deadlines were reviewed, with focus on upcoming deliverables and stakeholder commitments. Key risks and dependencies were identified that may impact program timeline and resource allocation.
        
        Next steps include addressing identified action items, monitoring risk mitigation efforts, and ensuring alignment across all program stakeholders for successful delivery.
        """
        
        return summary.strip()

    def generate_agile_tpm_minutes(self, transcription_text: str, meeting_title: str = "TPM Meeting Minutes", 
                                 meeting_date: Optional[datetime] = None) -> Dict[str, Any]:
        """Generate complete Agile TPM meeting minutes."""
        if meeting_date is None:
            meeting_date = datetime.now()
        
        # Parse transcription
        parsed_content = self.parse_transcription(transcription_text)
        
        # Extract Agile artifacts
        artifacts = self.extract_agile_artifacts(parsed_content)
        
        # Extract sprint/timeline info
        sprint_info = self.extract_sprint_info(parsed_content)
        
        # Generate TPM summary
        tpm_summary = self.generate_tpm_summary(parsed_content, artifacts, sprint_info)
        
        # Get unique speakers
        speakers = list(set([item['speaker'] for item in parsed_content]))
        
        return {
            'title': meeting_title,
            'date': meeting_date,
            'speakers': speakers,
            'tpm_summary': tpm_summary,
            'artifacts': artifacts,
            'sprint_info': sprint_info,
            'transcription': parsed_content
        }

    def save_docx(self, meeting_data: Dict[str, Any], output_path: str) -> bool:
        """Save meeting minutes as DOCX document."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot save DOCX.")
            return False
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(meeting_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Meeting info
            doc.add_heading('Meeting Information', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = 'Date'
            info_table.cell(0, 1).text = meeting_data['date'].strftime('%B %d, %Y at %I:%M %p')
            info_table.cell(1, 0).text = 'Participants'
            info_table.cell(1, 1).text = ', '.join(meeting_data['speakers'])
            info_table.cell(2, 0).text = 'Meeting Type'
            info_table.cell(2, 1).text = 'Agile TPM Program Review'
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(meeting_data['tpm_summary'])
            
            # Program Metrics
            doc.add_heading('Program Metrics', level=1)
            metrics_table = doc.add_table(rows=1, cols=5)
            metrics_table.style = 'Table Grid'
            
            headers = ['Epics', 'User Stories', 'Technical Tasks', 'Risks', 'Action Items']
            for i, header in enumerate(headers):
                metrics_table.cell(0, i).text = header
            
            # Add metrics data
            metrics_row = metrics_table.add_row()
            metrics_row.cells[0].text = str(len(meeting_data['artifacts'].get('epics', [])))
            metrics_row.cells[1].text = str(len(meeting_data['artifacts'].get('user_stories', [])))
            metrics_row.cells[2].text = str(len(meeting_data['artifacts'].get('technical_tasks', [])))
            metrics_row.cells[3].text = str(len(meeting_data['artifacts'].get('risks', [])))
            metrics_row.cells[4].text = str(len(meeting_data['artifacts'].get('action_items', [])))
            
            # Epics
            if meeting_data['artifacts'].get('epics'):
                doc.add_heading('Epics & Major Initiatives', level=1)
                for epic in meeting_data['artifacts']['epics']:
                    doc.add_heading(epic['title'], level=2)
                    doc.add_paragraph(epic['description'])
                    doc.add_paragraph(f"Priority: {epic['priority']}")
            
            # Risks
            if meeting_data['artifacts'].get('risks'):
                doc.add_heading('Risks & Issues', level=1)
                for risk in meeting_data['artifacts']['risks']:
                    doc.add_heading(risk['title'], level=2)
                    doc.add_paragraph(risk['description'])
                    doc.add_paragraph(f"Impact: {risk['impact']} | Probability: {risk['probability']}")
            
            # Decisions
            if meeting_data['artifacts'].get('decisions'):
                doc.add_heading('Decisions Made', level=1)
                for decision in meeting_data['artifacts']['decisions']:
                    doc.add_heading(decision['title'], level=2)
                    doc.add_paragraph(decision['description'])
                    doc.add_paragraph(f"Rationale: {decision['rationale']}")
            
            # Action Items
            if meeting_data['artifacts'].get('action_items'):
                doc.add_heading('Action Items', level=1)
                for item in meeting_data['artifacts']['action_items']:
                    doc.add_heading(f"Action: {item['action'][:50]}...", level=2)
                    doc.add_paragraph(f"Owner: {item['owner']}")
                    doc.add_paragraph(f"Due Date: {item['due_date']}")
                    doc.add_paragraph(f"Priority: {item['priority']}")
            
            # Timeline
            if meeting_data['sprint_info'].get('dates') or meeting_data['sprint_info'].get('sprints'):
                doc.add_heading('Timeline & Milestones', level=1)
                if meeting_data['sprint_info'].get('dates'):
                    doc.add_heading('Key Dates', level=2)
                    for date in meeting_data['sprint_info']['dates']:
                        doc.add_paragraph(f"• {date}")
                if meeting_data['sprint_info'].get('sprints'):
                    doc.add_heading('Sprint References', level=2)
                    for sprint in meeting_data['sprint_info']['sprints']:
                        doc.add_paragraph(f"• {sprint}")
            
            # Save document
            doc.save(output_path)
            logger.info(f"DOCX document saved to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to save DOCX: {e}")
            return False

    def save_html(self, meeting_data: Dict[str, Any], output_path: str) -> bool:
        """Save meeting minutes as HTML document."""
        if not JINJA2_AVAILABLE:
            return self._save_html_simple(meeting_data, output_path)
        
        try:
            template_str = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; max-width: 1000px; margin: 0 auto; padding: 20px; background-color: #f8f9fa; }
        .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 20px; }
        .section { background: white; margin-bottom: 20px; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
        .section h2 { color: #2c3e50; border-left: 4px solid #3498db; padding-left: 15px; margin-top: 0; }
        .epic { background: #e8f5e8; border-left: 4px solid #27ae60; padding: 10px; margin: 5px 0; }
        .risk { background: #ffebee; border-left: 4px solid #f44336; padding: 10px; margin: 5px 0; }
        .decision { background: #e0f2f1; border-left: 4px solid #009688; padding: 10px; margin: 5px 0; }
        .action-item { background: #fff8e1; border-left: 4px solid #ffc107; padding: 10px; margin: 5px 0; }
        .transcription { max-height: 400px; overflow-y: auto; border: 1px solid #ddd; padding: 15px; background: #fafafa; }
        .speaker { font-weight: bold; color: #3498db; }
        .timestamp { color: #7f8c8d; font-size: 0.9em; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ title }}</h1>
        <p><strong>Date:</strong> {{ date.strftime('%B %d, %Y at %I:%M %p') }}</p>
        <p><strong>Participants:</strong> {{ speakers | join(', ') }}</p>
        <p><strong>Meeting Type:</strong> Agile TPM Program Review</p>
    </div>

    <div class="section">
        <h2>📋 Executive Summary</h2>
        <p>{{ tpm_summary }}</p>
    </div>

    <div class="section">
        <h2>🎯 Epics & Major Initiatives</h2>
        {% for epic in artifacts.epics %}
        <div class="epic">
            <strong>{{ epic.title }}</strong><br>
            {{ epic.description }}<br>
            <small>Priority: {{ epic.priority }}</small>
        </div>
        {% endfor %}
    </div>

    <div class="section">
        <h2>⚠️ Risks & Issues</h2>
        {% for risk in artifacts.risks %}
        <div class="risk">
            <strong>{{ risk.title }}</strong><br>
            {{ risk.description }}<br>
            <small>Impact: {{ risk.impact }} | Probability: {{ risk.probability }}</small>
        </div>
        {% endfor %}
    </div>

    <div class="section">
        <h2>✅ Decisions Made</h2>
        {% for decision in artifacts.decisions %}
        <div class="decision">
            <strong>{{ decision.title }}</strong><br>
            {{ decision.description }}<br>
            <small>Rationale: {{ decision.rationale }}</small>
        </div>
        {% endfor %}
    </div>

    <div class="section">
        <h2>📝 Action Items</h2>
        {% for item in artifacts.action_items %}
        <div class="action-item">
            <strong>Action:</strong> {{ item.action }}<br>
            <strong>Owner:</strong> {{ item.owner }}<br>
            <strong>Due Date:</strong> {{ item.due_date }}<br>
            <small>Priority: {{ item.priority }}</small>
        </div>
        {% endfor %}
    </div>

    <div class="section">
        <h2>💬 Full Transcription</h2>
        <div class="transcription">
            {% for item in transcription %}
            <p>
                {% if item.timestamp %}<span class="timestamp">[{{ item.timestamp }}]</span> {% endif %}
                <span class="speaker">{{ item.speaker }}:</span> {{ item.content }}
            </p>
            {% endfor %}
        </div>
    </div>
</body>
</html>
            """
            
            template = Template(template_str)
            html_content = template.render(**meeting_data)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML document saved to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to save HTML: {e}")
            return False

    def _save_html_simple(self, meeting_data: Dict[str, Any], output_path: str) -> bool:
        """Simple HTML generation without Jinja2."""
        try:
            artifacts = meeting_data['artifacts']
            
            html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{meeting_data['title']}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; max-width: 1000px; margin: 0 auto; padding: 20px; background-color: #f8f9fa; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
        .section {{ background: white; margin-bottom: 20px; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .section h2 {{ color: #2c3e50; border-left: 4px solid #3498db; padding-left: 15px; margin-top: 0; }}
        .epic {{ background: #e8f5e8; border-left: 4px solid #27ae60; padding: 10px; margin: 5px 0; }}
        .risk {{ background: #ffebee; border-left: 4px solid #f44336; padding: 10px; margin: 5px 0; }}
        .decision {{ background: #e0f2f1; border-left: 4px solid #009688; padding: 10px; margin: 5px 0; }}
        .action-item {{ background: #fff8e1; border-left: 4px solid #ffc107; padding: 10px; margin: 5px 0; }}
        .transcription {{ max-height: 400px; overflow-y: auto; border: 1px solid #ddd; padding: 15px; background: #fafafa; }}
        .speaker {{ font-weight: bold; color: #3498db; }}
        .timestamp {{ color: #7f8c8d; font-size: 0.9em; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{meeting_data['title']}</h1>
        <p><strong>Date:</strong> {meeting_data['date'].strftime('%B %d, %Y at %I:%M %p')}</p>
        <p><strong>Participants:</strong> {', '.join(meeting_data['speakers'])}</p>
        <p><strong>Meeting Type:</strong> Agile TPM Program Review</p>
    </div>

    <div class="section">
        <h2>📋 Executive Summary</h2>
        <p>{meeting_data['tpm_summary']}</p>
    </div>

    <div class="section">
        <h2>🎯 Epics & Major Initiatives</h2>
        {''.join([f'<div class="epic"><strong>{epic["title"]}</strong><br>{epic["description"]}<br><small>Priority: {epic["priority"]}</small></div>' for epic in artifacts.get('epics', [])])}
    </div>

    <div class="section">
        <h2>⚠️ Risks & Issues</h2>
        {''.join([f'<div class="risk"><strong>{risk["title"]}</strong><br>{risk["description"]}<br><small>Impact: {risk["impact"]} | Probability: {risk["probability"]}</small></div>' for risk in artifacts.get('risks', [])])}
    </div>

    <div class="section">
        <h2>✅ Decisions Made</h2>
        {''.join([f'<div class="decision"><strong>{decision["title"]}</strong><br>{decision["description"]}<br><small>Rationale: {decision["rationale"]}</small></div>' for decision in artifacts.get('decisions', [])])}
    </div>

    <div class="section">
        <h2>📝 Action Items</h2>
        {''.join([f'<div class="action-item"><strong>Action:</strong> {item["action"]}<br><strong>Owner:</strong> {item["owner"]}<br><strong>Due Date:</strong> {item["due_date"]}<br><small>Priority: {item["priority"]}</small></div>' for item in artifacts.get('action_items', [])])}
    </div>

    <div class="section">
        <h2>💬 Full Transcription</h2>
        <div class="transcription">
            {''.join([f'<p><span class="speaker">{item["speaker"]}:</span> {item["content"]}</p>' for item in meeting_data['transcription']])}
        </div>
    </div>
</body>
</html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML document saved to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to save HTML: {e}")
            return False


# Flask API Service
if FLASK_AVAILABLE:
    app = Flask(__name__)
    meeting_minutes_service = MeetingMinutesService()

    @app.route('/health', methods=['GET'])
    def health_check():
        """Health check endpoint."""
        return jsonify({
            "status": "healthy",
            "service": "meeting_minutes",
            "openai_available": OPENAI_AVAILABLE,
            "jinja2_available": JINJA2_AVAILABLE,
            "docx_available": DOCX_AVAILABLE
        })

    @app.route('/generate-minutes', methods=['POST'])
    def generate_minutes():
        """Generate meeting minutes from transcription."""
        try:
            data = request.get_json()
            
            if not data or 'transcription_text' not in data:
                return jsonify({
                    "success": False,
                    "error": "transcription_text is required"
                }), 400
            
            transcription_text = data['transcription_text']
            meeting_title = data.get('meeting_title', 'TPM Meeting Minutes')
            meeting_date = data.get('meeting_date')
            
            if meeting_date:
                try:
                    meeting_date = datetime.fromisoformat(meeting_date)
                except ValueError:
                    meeting_date = None
            
            result = meeting_minutes_service.generate_agile_tpm_minutes(
                transcription_text, meeting_title, meeting_date
            )
            
            return jsonify({
                "success": True,
                "data": result
            })
            
        except Exception as e:
            logger.error(f"API error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/save-docx', methods=['POST'])
    def save_docx():
        """Save meeting minutes as DOCX."""
        try:
            data = request.get_json()
            
            if not data or 'meeting_data' not in data or 'output_path' not in data:
                return jsonify({
                    "success": False,
                    "error": "meeting_data and output_path are required"
                }), 400
            
            success = meeting_minutes_service.save_docx(
                data['meeting_data'], data['output_path']
            )
            
            return jsonify({
                "success": success,
                "error": None if success else "Failed to save DOCX"
            })
            
        except Exception as e:
            logger.error(f"DOCX save error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    @app.route('/save-html', methods=['POST'])
    def save_html():
        """Save meeting minutes as HTML."""
        try:
            data = request.get_json()
            
            if not data or 'meeting_data' not in data or 'output_path' not in data:
                return jsonify({
                    "success": False,
                    "error": "meeting_data and output_path are required"
                }), 400
            
            success = meeting_minutes_service.save_html(
                data['meeting_data'], data['output_path']
            )
            
            return jsonify({
                "success": success,
                "error": None if success else "Failed to save HTML"
            })
            
        except Exception as e:
            logger.error(f"HTML save error: {e}")
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500

    def run_service(host='localhost', port=5002, debug=False, api_key=None):
        """Run the meeting minutes service."""
        if api_key:
            meeting_minutes_service.api_key = api_key
            if OPENAI_AVAILABLE:
                meeting_minutes_service.client = openai.OpenAI(api_key=api_key)
        
        logger.info(f"Starting Meeting Minutes Service on {host}:{port}")
        app.run(host=host, port=port, debug=debug)

else:
    def run_service(host='localhost', port=5002, debug=False, api_key=None):
        """Fallback when Flask is not available."""
        logger.error("Flask not available. Cannot run as web service.")
        sys.exit(1)


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Meeting Minutes Generation Microservice")
    parser.add_argument('--host', default='localhost', help='Host to bind to')
    parser.add_argument('--port', type=int, default=5002, help='Port to bind to')
    parser.add_argument('--api-key', help='OpenAI API key for AI features')
    parser.add_argument('--debug', action='store_true', help='Enable debug mode')
    
    args = parser.parse_args()
    
    run_service(host=args.host, port=args.port, debug=args.debug, api_key=args.api_key)
